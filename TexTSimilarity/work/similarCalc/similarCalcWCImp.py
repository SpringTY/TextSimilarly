# 输入设定文档目录
# 输出相似度分数矩阵和相似度超过阈值的文档对

import os
import docx
import jieba

# pdfresourcemanager：PDF资源管理器：用于存储共享资源，如字体或图像
# LAParams：PDF参数分析器：分析pdf文件参数

filterLength = 10  # 过滤文本长度值


# 重命名所有报告名称为学号.文件格式

def getDocx(path: str):
    '''
    提取docx文档中的文本内容
    '''
    # with docx.Document(path) as f:
    f = docx.Document(path)
    return (''.join(filterRow([p.text for p in f.paragraphs])))


def filterRow(doc: list):
    '''
    过滤文本中不需要的内容，可以添加更多规则
    '''
    for para in list(reversed(doc)):
        if '\xa0' in para:  # 删去代码段，有的高亮部分可以删去，有的删不去
            doc.remove(para)
        elif para in ['\n', '', ' ']:  # 删去没有文字的段
            doc.remove(para)
        elif len(para) < filterLength:  # 长度小于十，删去
            doc.remove(para)
    return doc


# 从路径中读取文章
def readDocFromPath(path: str):
    docNameList, docTextList = [], []  # 文件名列表、文件内容列表
    failDocList = []
    for root, dirs, files in os.walk(path):
        '''
        os.walk是用于通过在目录树中游走输出在目录中的文件名，向上或者向下；
        root 所指的是当前正在遍历的这个文件夹的本身的地址
        dirs 是一个 list ，内容是该文件夹中所有的目录的名字(不包括子目录)
        files 同样是 list , 内容是该文件夹中所有的文件名字（str）(不包括子目录)
        '''
        for file in files:  # 对文件进行识别和分类，不同文件做不同处理
            try:
                extension = file.split('.')[-1]
                if extension == 'docx':  # 如果扩展名是docx
                    docTextList.append(getDocx(os.path.join(root, file)))
                    docNameList.append(file)
            except:
                failDocList.append(os.path.join(root, file))
    print('以下文件读取失败')
    print('\n'.join(failDocList))
    print('文件读取完成')
    return (docNameList, docTextList)


# 去标点（在求分词结果时用）
def remove_list():
    remove_list = ['，', '：', '。', '《', '》', '\n', '_', '“', '”', '、', ' ', ':', '（', '）', '.', '(', ')', '=']
    return remove_list


# 分词结果
def jieba_result(content):
    res = jieba.lcut(content)
    for i in remove_list():
        while i in res:
            res.remove(i)
    return res


# 算词频
def top_n(res, res_all):
    dic = {}
    for i in res_all:
        if i not in res:
            dic[i] = 0
        else:
            dic[i] = res.count(i)
    top_n_res = []
    for i in dic.items():
        top_n_res.append(i[1])
    # print(top_n_res)
    return top_n_res


# 计算相似度
def CalculateCos(gwyList, subList):
    gwyLen = 0
    for gwynum in gwyList:
        gwyLen = gwyLen + gwynum ** 2
    gwyLen = gwyLen ** 0.5
    subLen = 0
    for sub in subList:
        subLen = subLen + sub ** 2
    subLen = subLen ** 0.5
    # return subLen
    totalLen = len(gwyList)
    fenmu = 0
    for i in range(0, totalLen):
        fenmu = fenmu + subList[i] * gwyList[i]
    # print(fenmu / (subLen * gwyLen))
    return fenmu / (subLen * gwyLen)


def similarCalcWCImp(cmsFilePath):
    testPath = cmsFilePath
    docNameList, docTextList = readDocFromPath(testPath)
    # 不使用tfidf
    merge = "".join(docTextList)  # 合并所有文本
    merge_Participle = jieba_result(merge)  # 所有文本做分词结果（做分词慢）
    top_n_res = []  # 词频结果
    for i in docTextList:
        top_n_res.append(top_n(jieba_result(i), merge_Participle))  # 算词频
    similarity = []
    for i in range(len(top_n_res)):
        for j in range(i + 1, len(top_n_res)):
            result = CalculateCos(top_n_res[i], top_n_res[j])
            if (result >= 0.9):
                print("result >= 0.9", docNameList[i], "与", docNameList[j], "相似")
            similarity.append({
                'from': i,
                'to': j,
                'similarity': result
            })
    return docNameList, similarity


if __name__ == '__main__':
    docNameList, similarity = similarCalcWCImp(
        'C:/Users/Administrator/Documents/workspace/python/TextSimilarly/upload_files/test')
